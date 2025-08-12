from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
import os
import re

class DocumentGenerator:
    def __init__(self, output_dir='generated'):
        self.output_dir = output_dir
        self.ensure_output_dir()
    
    def ensure_output_dir(self):
        """Crea el directorio de salida si no existe"""
        os.makedirs(self.output_dir, exist_ok=True)
    
    def get_month_name(self, month_number):
        """Convierte número de mes a nombre en español"""
        months = [
            '', 'enero', 'febrero', 'marzo', 'abril', 'mayo', 'junio',
            'julio', 'agosto', 'septiembre', 'octubre', 'noviembre', 'diciembre'
        ]
        return months[month_number]
    
    def generate_resolution(self, aprendiz_data, plantilla_data, numero_resolucion):
        """Genera una resolución en formato Word según estándar SENA"""
        
        # Crear documento
        doc = Document()
        
        # Configurar márgenes y estilo
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1.2)
            section.right_margin = Inches(1)
        
        # Agregar encabezado con logos (simulado con texto)
        header_para = doc.add_paragraph()
        header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        header_run = header_para.add_run("SERVICIO NACIONAL DE APRENDIZAJE - SENA")
        header_run.bold = True
        header_run.font.size = 14
        
        # Regional y Centro
        centro_para = doc.add_paragraph()
        centro_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        centro_run = centro_para.add_run("REGIONAL BOYACÁ - CENTRO MINERO")
        centro_run.bold = True
        centro_run.font.size = 12
        
        # Agregar línea en blanco
        doc.add_paragraph()
        
        # Preparar datos para reemplazo
        fecha_actual = datetime.now()
        
        datos_reemplazo = {
            'numero_resolucion': numero_resolucion,
            'nombres': aprendiz_data['nombres'],
            'apellidos': aprendiz_data['apellidos'],
            'tipo_documento': aprendiz_data['tipo_documento'],
            'numero_documento': aprendiz_data['numero_documento'],
            'programa': aprendiz_data['programa'],
            'ficha': aprendiz_data['ficha'],
            'ciudad': 'Sogamoso',
            'dia': str(fecha_actual.day),
            'mes': self.get_month_name(fecha_actual.month),
            'año': str(fecha_actual.year),
            'fecha_completa': f"{fecha_actual.day} de {self.get_month_name(fecha_actual.month)} de {fecha_actual.year}"
        }
        
        # Título de la resolución
        titulo_para = doc.add_paragraph()
        titulo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        titulo_run = titulo_para.add_run(f"RESOLUCIÓN No. {numero_resolucion} DE {fecha_actual.year}")
        titulo_run.bold = True
        titulo_run.font.size = 14
        
        # Subtítulo según tipo de resolución
        if plantilla_data['tipo'] == 'APOYO_SOSTENIMIENTO':
            subtitulo = f'Por la cual se otorga apoyo de sostenimiento al aprendiz {datos_reemplazo["nombres"]} {datos_reemplazo["apellidos"]}'
        elif plantilla_data['tipo'] == 'TRANSPORTE':
            subtitulo = f'Por la cual se otorga apoyo de transporte al aprendiz {datos_reemplazo["nombres"]} {datos_reemplazo["apellidos"]}'
        elif plantilla_data['tipo'] == 'MONITORIA':
            subtitulo = f'Por la cual se designa como monitor académico al aprendiz {datos_reemplazo["nombres"]} {datos_reemplazo["apellidos"]}'
        else:
            subtitulo = plantilla_data.get('descripcion', 'Resolución administrativa')
            
        subtitulo_para = doc.add_paragraph()
        subtitulo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitulo_run = subtitulo_para.add_run(subtitulo)
        subtitulo_run.font.size = 11
        
        # Línea de código de formato
        doc.add_paragraph()
        codigo_para = doc.add_paragraph()
        codigo_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        codigo_run = codigo_para.add_run("GD-F-010 V05 Pag # 1")
        codigo_run.font.size = 10
        
        # Autoridad que expide
        autoridad_para = doc.add_paragraph()
        autoridad_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        autoridad_run = autoridad_para.add_run("EL SUBDIRECTOR (E) DEL CENTRO MINERO DEL SERVICIO NACIONAL DE APRENDIZAJE – SENA")
        autoridad_run.bold = True
        autoridad_run.font.size = 12
        
        # Fundamento legal
        doc.add_paragraph()
        fundamento_para = doc.add_paragraph()
        fundamento_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        fundamento_text = "En uso de sus facultades legales y reglamentarias, en especial las conferidas por los numerales 29 y 32 del artículo 27° del Decreto 249 de 2004 y el artículo 1 de la Resolución 00621 de 2013 y las conferidas por el director general de la Entidad, mediante Resolución 1-00618- del 17 de abril del año 2023, Acta de posesión No. 120 del 17 de abril de 2023."
        fundamento_para.add_run(fundamento_text)
        
        # CONSIDERANDO
        doc.add_paragraph()
        considerando_para = doc.add_paragraph()
        considerando_run = considerando_para.add_run("CONSIDERANDO:")
        considerando_run.bold = True
        considerando_run.font.size = 12
        
        # Agregar considerandos específicos según el tipo
        self._agregar_considerandos(doc, plantilla_data['tipo'])
        
        # RESUELVE
        doc.add_paragraph()
        resuelve_para = doc.add_paragraph()
        resuelve_run = resuelve_para.add_run("RESUELVE:")
        resuelve_run.bold = True
        resuelve_run.font.size = 12
        
        # Procesar contenido de la plantilla y reemplazar variables
        contenido = plantilla_data['contenido']
        for variable, valor in datos_reemplazo.items():
            contenido = contenido.replace(f'{{{variable}}}', str(valor))
        
        # Agregar artículos
        self._agregar_articulos(doc, contenido, datos_reemplazo)
        
        # Pie de resolución
        self._agregar_pie_resolucion(doc, datos_reemplazo)
        
        # Generar nombre de archivo
        filename = f"resolucion_{numero_resolucion.replace('-', '_')}_{aprendiz_data['numero_documento']}.docx"
        filepath = os.path.join(self.output_dir, filename)
        
        # Guardar documento
        doc.save(filepath)
        
        return filepath
    
    def _agregar_considerandos(self, doc, tipo_resolucion):
        """Agrega los considerandos específicos según el tipo de resolución"""
        
        considerandos_base = [
            'Que el artículo 6º del Decreto 2375 de 1974 establece la creación del Fondo Nacional de Formación Profesional de la Industria de la Construcción.',
            'Que el SENA tiene la responsabilidad de administrar los recursos destinados al bienestar de los aprendices.',
            'Que es necesario garantizar el apoyo a los aprendices durante su proceso de formación.'
        ]
        
        # Agregar considerandos específicos según el tipo
        if tipo_resolucion == 'APOYO_SOSTENIMIENTO':
            considerandos_especificos = [
                'Que el aprendiz cumple con los requisitos establecidos para el otorgamiento del apoyo de sostenimiento.',
                'Que existe disponibilidad presupuestal para atender la solicitud.'
            ]
        elif tipo_resolucion == 'TRANSPORTE':
            considerandos_especificos = [
                'Que se requiere facilitar el desplazamiento del aprendiz hacia el centro de formación.',
                'Que el apoyo de transporte contribuye a la permanencia en el programa formativo.'
            ]
        elif tipo_resolucion == 'MONITORIA':
            considerandos_especificos = [
                'Que el aprendiz ha demostrado excelencia académica y competencias para ejercer monitoria.',
                'Que la monitoria académica fortalece el proceso de formación integral.'
            ]
        else:
            considerandos_especificos = []
        
        # Agregar todos los considerandos al documento
        for considerando in considerandos_base + considerandos_especificos:
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            para.add_run(considerando)
            para.space_after = 6
    
    def _agregar_articulos(self, doc, contenido, datos_reemplazo):
        """Agrega los artículos de la resolución"""
        
        # Dividir contenido en artículos
        articulos = contenido.split('ARTÍCULO')
        
        for i, articulo in enumerate(articulos):
            if articulo.strip():
                art_para = doc.add_paragraph()
                if i == 0:
                    # Primer artículo
                    art_run = art_para.add_run("ARTÍCULO " + articulo.strip())
                else:
                    art_run = art_para.add_run("ARTÍCULO" + articulo.strip())
                
                # Hacer negrita la palabra ARTÍCULO y el número
                art_text = art_run.text
                if ":" in art_text:
                    parts = art_text.split(":", 1)
                    art_para.clear()
                    bold_part = art_para.add_run(parts[0] + ":")
                    bold_part.bold = True
                    art_para.add_run(" " + parts[1].strip())
                
                art_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                art_para.space_after = 12
    
    def _agregar_pie_resolucion(self, doc, datos_reemplazo):
        """Agrega el pie de la resolución con firmas y datos"""
        
        # COMUNÍQUESE Y CÚMPLASE
        doc.add_paragraph()
        comuniquese_para = doc.add_paragraph()
        comuniquese_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        comuniquese_run = comuniquese_para.add_run("COMUNÍQUESE Y CÚMPLASE")
        comuniquese_run.bold = True
        comuniquese_run.font.size = 12
        
        # Lugar y fecha
        fecha_para = doc.add_paragraph()
        fecha_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fecha_para.add_run(f'Dado en {datos_reemplazo["ciudad"]}, a los {datos_reemplazo["dia"]} días del mes de {datos_reemplazo["mes"]} de {datos_reemplazo["año"]}')
        
        # Espacio para firma
        doc.add_paragraph("\n\n")
        
        # Línea de firma
        firma_para = doc.add_paragraph()
        firma_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        firma_para.add_run("_" * 50)
        
        # Nombre del firmante
        nombre_para = doc.add_paragraph()
        nombre_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        nombre_run = nombre_para.add_run("Harvey Yadiver Dimaté Rodríguez")
        nombre_run.bold = True
        
        # Cargo
        cargo_para = doc.add_paragraph()
        cargo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cargo_para.add_run("Subdirector (E) Centro Minero")
        
        # Revisiones
        doc.add_paragraph("\n")
        revisiones = [
            "VoBo: Julieth Alejandra Viancha Torres: Jurídica Subdirección.",
            "Revisó: Blanca Katherin Gómez Viancha – Coordinadora de Formación.",
            "Revisó: Eliana Cruz Mora - Líder de Bienestar.",
            "Elaboró: Claudia Patricia Rincón Vija - Apoyo socioeconómico."
        ]
        
        for revision in revisiones:
            rev_para = doc.add_paragraph()
            rev_para.add_run(revision)
            rev_para.space_after = 6
    
    def create_batch_summary(self, generated_files):
        """Crea un resumen de generación masiva"""
        doc = Document()
        
        # Título
        title = doc.add_heading('Resumen de Generación de Resoluciones', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Información general
        fecha_generacion = datetime.now().strftime('%d/%m/%Y %H:%M:%S')
        doc.add_paragraph(f'Fecha de generación: {fecha_generacion}')
        
        total_files = len(generated_files)
        successful_files = len([f for f in generated_files if f['status'] == 'success'])
        failed_files = total_files - successful_files
        
        doc.add_paragraph(f'Total de resoluciones: {total_files}')
        doc.add_paragraph(f'Generadas exitosamente: {successful_files}')
        doc.add_paragraph(f'Fallidas: {failed_files}')
        
        # Tabla de resultados
        doc.add_heading('Detalle de Resoluciones', level=1)
        
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        
        # Encabezados
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Aprendiz'
        hdr_cells[1].text = 'Documento'
        hdr_cells[2].text = 'No. Resolución'
        hdr_cells[3].text = 'Estado'
        
        # Datos
        for file_info in generated_files:
            row_cells = table.add_row().cells
            row_cells[0].text = file_info['aprendiz']
            row_cells[1].text = file_info['numero_documento']
            row_cells[2].text = file_info['numero_resolucion']
            row_cells[3].text = file_info['status'].upper()
        
        # Guardar resumen
        summary_filename = f"resumen_generacion_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        summary_filepath = os.path.join(self.output_dir, summary_filename)
        doc.save(summary_filepath)
        
        return summary_filepath
        